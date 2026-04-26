import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_title(doc, text):
    title = doc.add_heading(text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')

def add_chapter(doc, title, text_block):
    heading = doc.add_heading(title, level=1)
    heading.style.font.color.rgb = RGBColor(0, 51, 102)
    
    # Process structured text block
    lines = text_block.strip().split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('### '):
            h3 = doc.add_heading(line[4:].strip(), level=3)
            h3.style.font.color.rgb = RGBColor(0, 102, 204)
        elif line.startswith('## '):
            h2 = doc.add_heading(line[3:].strip(), level=2)
            h2.style.font.color.rgb = RGBColor(0, 76, 153)
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            p.style.font.name = 'Arial'
            p.style.font.size = Pt(11)
        else:
            p = doc.add_paragraph(line)
            p.style.font.name = 'Arial'
            p.style.font.size = Pt(11)

def create_guide():
    doc = Document()
    add_title(doc, "Guide de Formation Complet et Ultime\nAdministrateurs Serenity")
    
    intro = """
Ce guide exhaustif est conçu pour les équipes d'administration, de gestion financière et de conformité de Serenity. Il aborde de manière opérationnelle et fonctionnelle les rouages internes de la plateforme. Vous y trouverez les processus étape par étape, la mécanique des pénalités, et les filets de sécurité intégrés.
    """
    add_chapter(doc, "Introduction", intro)
    doc.add_page_break()

    c1 = """
Le tableau de bord (Dashboard) est votre salle de contrôle. Il se rafraîchit en temps réel pour présenter la santé financière globale du système.
## 1. Métriques Clés Générales
- Fonds en transit (Liquidités totales) : Affiche la somme des soldes présents sur les comptes courants et les cagnottes internes non décaissées.
- Créances (Impayés globaux) : Le total des pénalités et crédits non remboursés.
- Utilisateurs Actifs et Adhésions en cours.
## 2. Le Panneau des Alertes et Raccourcis
- Alertes KYC : Notifie immédiatement lorsqu'un nouvel utilisateur vient de soumettre ses documents d'identité.
- Demandes de Retrait : Regroupe les demandes de décaissement manuelles des gains de parrainage et des cagnottes.
- Alertes d'Intégrité : (Cf. Chapitre 10) Informe d'une corruption de la base de données.
## 3. Flux Opérationnel Quotidien
L'officier de crédit ou l'administrateur doit :
1. Consulter les alertes KYC et filtrer les nouveaux.
2. Traiter les dossiers Nano-crédit en attente.
3. Vérifier que la tuile de Sécurité globale reste sur "Intègre".
    """
    add_chapter(doc, "Chapitre 1 : Tableau de bord et Synthèse", c1)

    c2 = """
La connaissance client (KYC) est la première ligne de défense contre le blanchiment et la fraude aux crédits.
## 1. Parcours de Soumission membre
Lors de l'inscription via l'app Serenity, le client n'a accès qu'au niveau 0 (Lecture seule). Pour commencer à cotiser ou demander un nano-crédit, il doit :
- Envoyer une photo de sa Pièce d'Identité valide.
- Prendre un Selfie.
- Fournir ses informations géolocalisées.
## 2. Espace de Validation Admin
L'administrateur, depuis l'onglet "Clients > KYC en attente", examine visuellement la cohérence.
- Rejet : En cas de flou, de pièce invalide ou de falsification soupçonnée, le bouton "Rejeter" demande explicitement d'entrer un motif (ex: "ID illisible"). Le client reçoit aussitôt un SMS/Post Push l'invitant à reprendre.
- Validation : En cliquant sur "Valider", le membre passe au Statut Vérifié. Ses limites de transferts augmentent immédiatement selon la politique des plafonds.
## 3. Fiche détaillée du Membre
La fiche client contient 3 éléments distincts :
- Le profil KYC.
- Le portefeuille et ses sous-comptes internes (Compte Courant, Compte Cagnottes, Dette Crédit).
- Son "Indice de Qualité Garant" : C'est le score absolu de fidélité. Ce score augmente à chaque tontine honorée avec assiduité et garantit l'accès aux prêts supérieurs.
    """
    add_chapter(doc, "Chapitre 2 : Gestion des Membres & Processus KYC", c2)

    c3 = """
Les Cagnottes (anciennement cotisations libres) permettent de lever et sécuriser des fonds avec une séparation de contexte stricte.
## 1. Cagnottes Privées
Conçues pour la cellule familiale, les associations (tontine d'amis) ou les événements privés (décès, mariage).
- Confidentialité : Masquées aux utilisateurs non invités. L'accès nécessite un lien ou ajout direct par le créateur de la cagnotte.
- Le créateur de la cagnotte nomme l'administrateur de l'évènement. 
- Validation des versements : Si un membre paye par Mobile Money, la validation est immédiate. S'il déclare payer en Cash/Main à main (Virement Hors-ligne), l'administrateur de l'évènement OU le Super Admin mondial (vous) doit valider la transaction pour débloquer le crédit fictif sur la plateforme.
## 2. Cagnottes Publiques
Conçues pour le Crowdfunding, le soutien public national ou les appels aux dons médiatisés.
- Visibilité : La cagnotte apparait sur la page d'accueil de l'app mobile Serenity pour inciter les dons spontanés.
- Modération : Toute création de cagnotte publique par un usager déclenche le statut "En attente d'approbation". Un Admin Serenity DOIT comprendre l'objet de la levée et cliquer sur "Publier" pour qu'elle devienne accessible.
## 3. Retraits de fonds
Lorsque le maître de la cagnotte désire rompre la caisse (Décaissement total ou partiel) :
- Une demande de décaissement est envoyée aux administrateurs.
- Vous utilisez l'onglet "Demandes de Retraits" -> Vérifiez que les fonds sont réels -> Appuyez sur le bouton "Débloquer" pour injecter sur son compte courant ou relayer vers Mobile Money.
    """
    add_chapter(doc, "Chapitre 3 : Cagnottes (Épargne Publique et Privée)", c3)

    c4 = """
Le système d'Épargne Automatique solidaire basé sur des cycles fermés.
## 1. Déploiement et Séance
- L'administration configure le montant du tour (ex: 50.000 FCFA), le nombre de parts et la fréquence (Hebdomadaire/Mensuel).
- Une Tontine démarre uniquement quand tous les créneaux sont remplis et le capital de démarrage validé.
## 2. Périodes de Levée automatiques
Le Cronjob Serenity s'exécute chaque jour. Si la date correspond à la séance de levée :
- Le système débite automatiquement les comptes courants des participants (ou mobile money intégré) au bénéfice du gagnant du tour.
- Prélèvement bloqué : Si un client n'a pas les fonds, le système tente une relance SMS. À expiration du délai de grâce, le système génère systématiquement une amende/pénalité Tontine.
## 3. Fidélité
Mener à son terme un tour de tontine propulse radicalement la visibilité financière d'un usager :
- L'algorithme centralisera cette donnée pour le calcul du "Score IA" pour un financement (voir Chap. 5).
    """
    add_chapter(doc, "Chapitre 4 : La Puissance des Tontines Automatiques", c4)

    c5 = """
Module phare de création de revenus et de financement à la demande pour vos membres.
## 1. Parcours de Demande
Le membre effectue la demande depuis son mobile. Il vise l'un des Paliers pré-configurés par l'Admin (Ex: Palier 1 - Rencontre -> 35.000 XOF).
Certains paliers exigent d'indiquer des "Garants Internes". 
## 2. L'Évaluation à double facteur (Unique)
- Score IA (0 à 3) : Une intelligence déterministe (Le `AiRiskEvaluationService`) analyse l'historique du membre, la véracité de ses emails, ses retards en tontine et l'impayé des garants. Un profil vert est évalué à 0.
- Score Humain (0 à 3) : Depuis la page de traitement, vous (l'Admin) donnez votre évaluation après appel téléphonique ou analyse physique du profil.
## 3. L'Auto-Octroi Magique
Lorsque vous cliquez sur "Réévaluer manuellement", le serveur fusionne les deux notes. Ainsi, si Score Total (IA + Humain) < 2 :
- Le module appelle l'API de PayDunya en arrière plan. 
- Il initie instantanément le transfert réel de fonds Mobile Money sur le téléphone du client en quelques millisecondes, sans intervention humaine supplémentaire. 
## 4. Gestion des Garants et Impayés
- Le garant "Gèle" une fraction de ses avoirs en acceptant la responsabilité.
- Pénalités automatiques : Le moteur Cron rajoute le montant de pénalité de retard par jour (ou par mois selon la config de palier).
- Si la dette échoue : L'Admin clique sur le bouton "Recouvrer sur les Garants". Le processus va forcer le prélèvement des garants validés et liquider instantanément la créance dans le "Compte des Impayés" Serenity de l'application.
    """
    add_chapter(doc, "Chapitre 5 : Nano-Crédit : Scoring IA et Octroi", c5)

    c6 = """
Ce système permet d'engranger et de compenser la viralité de la base client.
## 1. Mécanisme des Commissions
- Configuration Globale : Dans les paramètres, l'Admin fixe la récompense (ex: 1.500 FCFA forfaitaire, ou 5% du premier versement de tontine).
- Validité : Le filleul doit effectuer sa 1ère transaction (ex: faire un dépôt KYC d'au moins 20.000F) pour que le solde `parrainage_gains` du parrain se valide. Le système de délai de grâce évite les comptes frauduleux jetables.
## 2. Retrait des Gains Parrainage
Le parrain demande un déblocage de ses avoirs :
- La demande s'affiche dans "Demande de Retraits", section Parrainage.
- Vous examinez l'origine : Nombre de filleuls qualifiés vs fictifs.
- Un clic "Débourser" envoie le flux financier vers le Compte Courant du membre.
    """
    add_chapter(doc, "Chapitre 6 : Le Système de Parrainage et Rémunérations", c6)

    c7 = """
Traçabilité parfaite de chaque centime dans le système. Serenity maintient des Caisses miroirs pour chaque activité.
## 1. Les Sous-comptes Modélisés
Derrière chaque membre se cachent plusieurs caisses comptables virtuelles (Compte Courant, Dette Crédit, Caisse Tontines). Un décaissement ne "sort" pas l'argent dans le vide : il transfère la valeur de la caisse centrale (Trésor Nano-Crédit) vers la Dette Crédit, puis de Dette vers Compte Courant.
## 2. Dashboard Financier Global (Caisse Dashboard)
La section "Flux Financiers" illustre avec des Graphiques dynamiques l'évolution des mouvements de la caisse centrale. Les "Transactions E/S" détaillent chaque ligne de code de libellé avec une précision chirurgicale (Debit vs Crédit).
    """
    add_chapter(doc, "Chapitre 7 : Comptabilité Avancée et Flux", c7)

    c8 = """
Serenity gère les agrégateurs télécom. PayDunya est l'API mère de gestion.
## 1. Fonctionnement
Dans le fichier de service, l'intégration comprend les paiements Pull (Payer sa tontine) et Push (Octroi Nano credit automatique).
## 2. Panneau des Méthodes
Vous pouvez à l'aide des paramètres de paiement suspendre immédiatement une méthode (exp: Wave Sénégal hors ligne) afin d'empêcher les fausses facturations lors d'un crash réseau de l'opérateur en Afrique de l'Ouest.
    """
    add_chapter(doc, "Chapitre 8 : Méthodes de Paiements et API", c8)

    c9 = """
Outil direct et simple de relance de conversion.
- Newsletters : Le module de Communication permet l'envoi de mails customisés ou de notifications Push aux téléphones.
- Annonces marketing : Ces campagnes peuvent inciter (Taux de tontine bonifiés) à télécharger la nouvelle mise à jour ou prévenir d'une maintenance serveur prévue.
    """
    add_chapter(doc, "Chapitre 9 : Campagnes Marketing & Communication", c9)

    c10 = """
L'arme secrète financière Serenity : Le Forensic Engine et la protection Checksum.
## 1. La Signature MD5 Checksum
Chaque montant enregistré en base de données de Mouvements de Caisses est doublé d'une "Empreinte" cryptographique secrète lors de sa création via l'interface standard.
Si un administrateur malhonnête entre en base de données (`phpMyAdmin` ou requêtes SQL) et falsifie un montant "10.000 -> 100.000", alors `Montant ≠ Checksum`.
## 2. Surveillance Cron et Alertes
Toutes les 10 minutes, un CronJob balaye toutes les tables (NanoCredit, Cagnottes, Transactions). La pastille "Intégration & Audit" dans le menu s'allume en ALERTE rouge et génère un ticket de corruption.
## 3. Options de Remédiation (Pour l'Admin Chef)
Face à la corruption ciblée, vous pouvez analyser l'historique complet, identifier la table modifiée illégalement (Log Details) et opérer :
- RESTAURER : Reconstruit de force l'ancienne vraie valeur à partir des archives d'audit.
- SUSPENDRE : Gèle immédiatement le profil client pour empêcher une fuite externe de capitaux.
- ACCEPTER : Modifie cryptographiquement le Checksum à la nouvelle valeur pour lui redonner une intégrité normale (Utile s'il s'agissait d'une correction légitime effectuée techniquement).
    """
    add_chapter(doc, "Chapitre 10 : Sécurité Plénière, Checksums et Forensic", c10)

    c11 = """
Panneau de commandement exclusif à la racine de la plateforme.
## Configuration Dynamique
Les paramètres vous permettent d'éditer le coeur du système sans redéploiement informatique :
- Limites KYC (Ex: Seuil retrait avant KYC).
- Mots clés SMS ou numéros d'Assistance.
- Plafond d'engagement aux Nano-Crédits en fonction du rang.
- Taux de Récompenses fidélité Tontine.
Attention, chaque changement impactera des centaines d'utilisateurs actifs immédiatement. Une modification majeure est conseillée en heures creuses.
    """
    add_chapter(doc, "Chapitre 11 : Paramétrage Global Serenity", c11)

    # Save
    doc.save('Guide_Formation_Admin_Serenity_Detaillé.docx')
    print("FINISHED")

if __name__ == '__main__':
    create_guide()
